#!/usr/bin/env python3
"""Convert BRD markdown to .docx with proper formatting."""

import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INPUT_FILE = "/Users/karthik/Desktop/college application/BRD_CareerGuidance_Platform.md"
OUTPUT_FILE = "/Users/karthik/Desktop/college application/BRD_CareerGuidance_Platform.docx"


def set_cell_shading(cell, color):
    """Set background color for a table cell."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_borders(cell, color="AAAAAA"):
    """Set borders for a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for border_name in ["top", "left", "bottom", "right"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def format_inline(paragraph, text):
    """Handle bold (**text**) and inline code (`text`) formatting."""
    parts = re.split(r'(\*\*.*?\*\*|`[^`]+`)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.size = paragraph.runs[0].font.size if paragraph.runs else Pt(10)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        else:
            paragraph.add_run(part)


def add_formatted_paragraph(doc, text, style="Normal", bold=False, size=None, color=None, space_after=None):
    """Add a paragraph with formatting."""
    p = doc.add_paragraph(style=style)
    # Handle inline formatting
    if "**" in text or "`" in text:
        if size:
            dummy = p.add_run("")
            dummy.font.size = size
        format_inline(p, text)
    else:
        run = p.add_run(text)
        run.bold = bold
        if size:
            run.font.size = size
        if color:
            run.font.color.rgb = color
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p


def parse_table(lines, start_idx):
    """Parse a markdown table starting at start_idx. Return (rows, end_idx)."""
    rows = []
    i = start_idx
    while i < len(lines):
        line = lines[i].strip()
        if not line.startswith("|"):
            break
        # Skip separator rows
        if re.match(r'^\|[\s\-:|]+\|$', line):
            i += 1
            continue
        cells = [c.strip() for c in line.split("|")[1:-1]]
        rows.append(cells)
        i += 1
    return rows, i


def add_table(doc, rows):
    """Add a formatted table to the document."""
    if not rows:
        return
    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"

    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j >= num_cols:
                break
            cell = table.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)

            # Clean text
            clean = cell_text.strip()

            if i == 0:
                # Header row
                set_cell_shading(cell, "2B579A")
                run = p.add_run(re.sub(r'\*\*([^*]+)\*\*', r'\1', clean))
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            else:
                set_cell_borders(cell)
                if "**" in clean or "`" in clean:
                    format_inline(p, clean)
                    for run in p.runs:
                        if not run.font.size:
                            run.font.size = Pt(9)
                else:
                    run = p.add_run(clean)
                    run.font.size = Pt(9)

    doc.add_paragraph()  # spacing


def add_code_block(doc, code_lines):
    """Add a code block with monospace formatting."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    text = "\n".join(code_lines)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def convert():
    with open(INPUT_FILE, "r", encoding="utf-8") as f:
        content = f.read()

    lines = content.split("\n")
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)

    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    i = 0
    in_code_block = False
    code_lines = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            i += 1
            continue

        # Horizontal rules
        if stripped == "---":
            doc.add_paragraph("_" * 60)
            i += 1
            continue

        # Code blocks
        if stripped.startswith("```"):
            if in_code_block:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Headers
        if stripped.startswith("# ") and not stripped.startswith("## "):
            title = stripped[2:].strip()
            p = doc.add_heading(title, level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        if stripped.startswith("## "):
            title = stripped[3:].strip()
            doc.add_heading(title, level=1)
            i += 1
            continue

        if stripped.startswith("### "):
            title = stripped[4:].strip()
            doc.add_heading(title, level=2)
            i += 1
            continue

        if stripped.startswith("#### "):
            title = stripped[5:].strip()
            doc.add_heading(title, level=3)
            i += 1
            continue

        # Tables
        if stripped.startswith("|"):
            rows, end_idx = parse_table(lines, i)
            add_table(doc, rows)
            i = end_idx
            continue

        # Bullet points
        if stripped.startswith("- "):
            text = stripped[2:]
            p = doc.add_paragraph(style="List Bullet")
            if "**" in text or "`" in text:
                format_inline(p, text)
            else:
                p.add_run(text)
            i += 1
            continue

        # Numbered list
        numbered = re.match(r'^(\d+)\.\s+(.*)', stripped)
        if numbered:
            text = numbered.group(2)
            p = doc.add_paragraph(style="List Number")
            if "**" in text or "`" in text:
                format_inline(p, text)
            else:
                p.add_run(text)
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        if "**" in stripped or "`" in stripped:
            format_inline(p, stripped)
        else:
            p.add_run(stripped)
        i += 1

    doc.save(OUTPUT_FILE)
    print(f"✅ Document saved to: {OUTPUT_FILE}")


if __name__ == "__main__":
    convert()
